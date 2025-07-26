"""
Word document processor.

This module handles Word document processing.
"""

from typing import Any

import docx


class WordProcessor:
    """Processes Word documents."""

    def process(self, file_path: str) -> dict[str, Any]:
        """Process a Word file and extract its content."""
        try:
            doc = docx.Document(file_path)

            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            return {
                "text": text_content,
                "file_path": file_path,
                "paragraph_count": len(doc.paragraphs),
            }
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
