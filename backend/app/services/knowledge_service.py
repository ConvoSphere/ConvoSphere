"""
Knowledge base service for document management and RAG functionality.

This module provides services for uploading, processing, and searching documents
in the knowledge base for retrieval-augmented generation.
"""

import logging
import mimetypes
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, List, Optional, Dict

from sqlalchemy import and_, or_, func
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.knowledge import Document, DocumentChunk, SearchQuery, Tag, DocumentProcessingJob, DocumentStatus, DocumentType

from .ai_service import AIService
from .document_processor import document_processor
from .embedding_service import embedding_service
from .weaviate_service import WeaviateService
from app.core.database import get_db

logger = logging.getLogger(__name__)


class TagService:
    """Service for managing document tags."""
    
    def __init__(self, db: Session):
        self.db = db
    
    def normalize_tag_name(self, tag_name: str) -> str:
        """Normalize tag name (lowercase, trim, etc.)."""
        return tag_name.lower().strip()
    
    def get_or_create_tag(self, tag_name: str) -> Tag:
        """Get existing tag or create new one."""
        normalized_name = self.normalize_tag_name(tag_name)
        
        tag = self.db.query(Tag).filter(Tag.name == normalized_name).first()
        if not tag:
            tag = Tag(name=normalized_name)
            self.db.add(tag)
            self.db.flush()  # Get the ID
        
        return tag
    
    def get_tags(self, user_id: str, limit: int = 100) -> List[Tag]:
        """Get all tags for a user's documents."""
        return (
            self.db.query(Tag)
            .join(Document.tags)
            .filter(Document.user_id == user_id)
            .order_by(Tag.usage_count.desc(), Tag.name)
            .limit(limit)
            .all()
        )
    
    def search_tags(self, query: str, user_id: str, limit: int = 20) -> List[Tag]:
        """Search tags by name."""
        normalized_query = self.normalize_tag_name(query)
        return (
            self.db.query(Tag)
            .join(Document.tags)
            .filter(
                and_(
                    Document.user_id == user_id,
                    Tag.name.contains(normalized_query)
                )
            )
            .order_by(Tag.usage_count.desc(), Tag.name)
            .limit(limit)
            .all()
        )
    
    def delete_unused_tags(self) -> int:
        """Delete tags that are not used by any documents."""
        unused_tags = self.db.query(Tag).filter(Tag.usage_count == 0).all()
        count = len(unused_tags)
        for tag in unused_tags:
            self.db.delete(tag)
        self.db.commit()
        return count


class MetadataExtractor:
    """Service for extracting metadata from documents."""
    
    @staticmethod
    def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF files."""
        try:
            import PyPDF2
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract document info
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    metadata['author'] = info.get('/Author')
                    metadata['title'] = info.get('/Title')
                    metadata['subject'] = info.get('/Subject')
                    metadata['creator'] = info.get('/Creator')
                    metadata['producer'] = info.get('/Producer')
                    
                    # Extract creation date
                    creation_date = info.get('/CreationDate')
                    if creation_date:
                        try:
                            # PDF date format: D:YYYYMMDDHHmmSSOHH'mm'
                            year_str = creation_date[2:6]
                            metadata['year'] = int(year_str)
                        except (ValueError, IndexError):
                            pass
                
                # Count pages
                metadata['page_count'] = len(pdf_reader.pages)
                
            return metadata
        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata: {e}")
            return {}
    
    @staticmethod
    def extract_word_metadata(file_path: str) -> Dict[str, Any]:
        """Extract metadata from Word documents."""
        try:
            from docx import Document
            metadata = {}
            
            doc = Document(file_path)
            
            # Extract core properties
            core_props = doc.core_properties
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.created:
                metadata['year'] = core_props.created.year
            
            # Count paragraphs as rough word count
            word_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
            metadata['word_count'] = word_count
            
            return metadata
        except Exception as e:
            logger.warning(f"Failed to extract Word metadata: {e}")
            return {}
    
    @staticmethod
    def detect_language(text: str) -> Optional[str]:
        """Detect language of text content."""
        try:
            from langdetect import detect
            return detect(text)
        except Exception as e:
            logger.warning(f"Failed to detect language: {e}")
            return None


class KnowledgeService:
    """Service for managing knowledge base documents and search."""

    def __init__(self, db=None):
        self.db = db or get_db()
        self.weaviate_service = WeaviateService()
        self.ai_service = AIService()
        self.tag_service = TagService(self.db)
        self.metadata_extractor = MetadataExtractor()

        # Ensure upload directory exists
        self.upload_dir = Path(get_settings().UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    def create_document(
        self,
        user_id: str,
        title: str,
        file_name: str,
        file_content: bytes,
        description: str | None = None,
        tags: list[str] | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        """Create a new document and save it to storage."""
        try:
            # Generate unique filename
            file_id = str(uuid.uuid4())
            file_extension = Path(file_name).suffix
            file_type = file_extension.lower().lstrip(".")

            # Determine MIME type
            mime_type, _ = mimetypes.guess_type(file_name)

            # Save file to storage
            file_path = self.upload_dir / f"{file_id}{file_extension}"
            with open(file_path, "wb") as f:
                f.write(file_content)

            # Determine document type
            document_type = self._determine_document_type(file_type, mime_type)

            # Extract metadata from file
            extracted_metadata = self._extract_document_metadata(file_path, file_type)
            if metadata:
                extracted_metadata.update(metadata)

            # Create document record
            document = Document(
                id=uuid.UUID(file_id),
                user_id=uuid.UUID(user_id),
                title=title,
                description=description or "",
                file_name=file_name,
                file_path=str(file_path),
                file_type=file_type,
                file_size=len(file_content),
                mime_type=mime_type,
                document_type=document_type,
                **extracted_metadata
            )

            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)

            # Add tags if provided
            if tags:
                for tag_name in tags:
                    document.add_tag(tag_name, self.db)

            logger.info(f"Created document {document.id} for user {user_id}")
            return document

        except Exception as e:
            logger.error(f"Error creating document: {e}")
            self.db.rollback()
            raise

    def _determine_document_type(self, file_type: str, mime_type: str) -> str:
        """Determine document type based on file extension and MIME type."""
        type_mapping = {
            'pdf': DocumentType.PDF,
            'doc': DocumentType.DOCUMENT,
            'docx': DocumentType.DOCUMENT,
            'txt': DocumentType.TEXT,
            'md': DocumentType.TEXT,
            'xlsx': DocumentType.SPREADSHEET,
            'xls': DocumentType.SPREADSHEET,
            'csv': DocumentType.SPREADSHEET,
            'pptx': DocumentType.PRESENTATION,
            'ppt': DocumentType.PRESENTATION,
            'jpg': DocumentType.IMAGE,
            'jpeg': DocumentType.IMAGE,
            'png': DocumentType.IMAGE,
            'mp3': DocumentType.AUDIO,
            'wav': DocumentType.AUDIO,
            'mp4': DocumentType.VIDEO,
            'py': DocumentType.CODE,
            'js': DocumentType.CODE,
            'java': DocumentType.CODE,
        }
        
        return type_mapping.get(file_type, DocumentType.OTHER)

    def _extract_document_metadata(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Extract metadata from document file."""
        metadata = {}
        
        try:
            if file_type == 'pdf':
                metadata.update(self.metadata_extractor.extract_pdf_metadata(file_path))
            elif file_type in ['doc', 'docx']:
                metadata.update(self.metadata_extractor.extract_word_metadata(file_path))
            
            # Detect language if we have text content
            if file_type in ['txt', 'md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                    language = self.metadata_extractor.detect_language(text)
                    if language:
                        metadata['language'] = language
                        metadata['word_count'] = len(text.split())
                        metadata['character_count'] = len(text)
        
        except Exception as e:
            logger.warning(f"Failed to extract metadata from {file_path}: {e}")
        
        return metadata

    async def process_document(self, document_id: str) -> bool:
        """Process a document by extracting text and creating chunks."""
        try:
            document = (
                self.db.query(Document).filter(Document.id == document_id).first()
            )
            if not document:
                raise ValueError(f"Document {document_id} not found")

            # Update status to processing
            document.status = DocumentStatus.PROCESSING
            self.db.commit()

            # Read file content
            with open(document.file_path, "rb") as f:
                file_content = f.read()

            # Process document using document processor
            result = document_processor.process_document(
                file_content, document.file_name,
            )

            if not result["success"]:
                raise ValueError(
                    f"Document processing failed: {result.get('error', 'Unknown error')}",
                )

            # Create document chunks from processed chunks
            processed_chunks = result["chunks"]
            chunks = []

            for i, processed_chunk in enumerate(processed_chunks):
                chunk = DocumentChunk(
                    id=uuid.uuid4(),
                    document_id=document.id,
                    chunk_index=i,
                    content=processed_chunk["content"],
                    token_count=processed_chunk["token_count"],
                    chunk_size=len(processed_chunk["content"]),
                    chunk_type=processed_chunk.get("chunk_type", "text"),
                    page_number=processed_chunk.get("page_number"),
                    section_title=processed_chunk.get("section_title"),
                    table_id=processed_chunk.get("table_id"),
                    figure_id=processed_chunk.get("figure_id"),
                )

                # Add chunk metadata
                chunk.chunk_metadata = {
                    "start_word": processed_chunk.get("start_word"),
                    "end_word": processed_chunk.get("end_word"),
                }

                chunks.append(chunk)

            # Generate embeddings for chunks
            chunk_texts = [chunk.content for chunk in chunks]
            embeddings = await embedding_service.generate_embeddings(chunk_texts)

            # Add embeddings to chunks and store in Weaviate
            for i, chunk in enumerate(chunks):
                if i < len(embeddings) and embeddings[i]:
                    chunk.embedding = embeddings[i]
                    chunk.embedding_model = get_settings().default_embedding_model
                    chunk.embedding_created_at = datetime.utcnow()

                    # Enhanced metadata for Weaviate
                    weaviate_metadata = {
                        "title": document.title,
                        "file_type": document.file_type,
                        "document_type": document.document_type,
                        "chunk_index": chunk.chunk_index,
                        "chunk_type": chunk.chunk_type,
                        "user_id": str(document.user_id),
                        "processing_engine": result["metadata"].get(
                            "processing_engine", "traditional",
                        ),
                        "author": document.author,
                        "language": document.language,
                        "year": document.year,
                    }

                    # Add chunk-specific metadata
                    if chunk.page_number:
                        weaviate_metadata["page_number"] = chunk.page_number
                    if chunk.section_title:
                        weaviate_metadata["section_title"] = chunk.section_title

                    # Store in Weaviate
                    self.weaviate_service.add_document_chunk(
                        chunk_id=str(chunk.id),
                        document_id=str(document.id),
                        content=chunk.content,
                        embedding=embeddings[i],
                        metadata=weaviate_metadata,
                    )

            # Save chunks to database
            self.db.add_all(chunks)

            # Update document metadata
            document.processing_engine = result["metadata"].get("processing_engine", "traditional")
            document.processing_options = result["metadata"].get("processing_options", {})
            
            # Update content statistics
            document.page_count = result["metadata"].get("page_count")
            document.word_count = result["metadata"].get("word_count")
            document.character_count = result["metadata"].get("character_count")
            
            # Detect language if not already set
            if not document.language and chunks:
                sample_text = chunks[0].content[:1000]  # Use first 1000 chars
                document.language = self.metadata_extractor.detect_language(sample_text)

            document.status = DocumentStatus.PROCESSED
            document.processed_at = datetime.utcnow()

            self.db.commit()

            logger.info(
                f"Successfully processed document {document_id} with {len(chunks)} chunks",
            )
            return True

        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            # Update document status to error
            document = (
                self.db.query(Document).filter(Document.id == document_id).first()
            )
            if document:
                document.status = DocumentStatus.ERROR
                document.error_message = str(e)
                self.db.commit()
            return False

    def get_documents(
        self,
        user_id: str,
        skip: int = 0,
        limit: int = 100,
        status: str | None = None,
        document_type: str | None = None,
        author: str | None = None,
        year: int | None = None,
        language: str | None = None,
        tag_names: List[str] | None = None,
    ) -> tuple[list[Document], int]:
        """Get documents for a user with advanced filtering."""
        query = self.db.query(Document).filter(Document.user_id == user_id)

        if status:
            query = query.filter(Document.status == status)
        if document_type:
            query = query.filter(Document.document_type == document_type)
        if author:
            query = query.filter(Document.author.ilike(f"%{author}%"))
        if year:
            query = query.filter(Document.year == year)
        if language:
            query = query.filter(Document.language == language)
        if tag_names:
            for tag_name in tag_names:
                normalized_tag = self.tag_service.normalize_tag_name(tag_name)
                query = query.join(Document.tags).filter(Tag.name == normalized_tag)

        total = query.count()
        documents = query.offset(skip).limit(limit).all()

        return documents, total

    def get_document(self, document_id: str, user_id: str) -> Document | None:
        """Get a specific document by ID."""
        return (
            self.db.query(Document)
            .filter(
                and_(
                    Document.id == document_id,
                    Document.user_id == user_id,
                ),
            )
            .first()
        )

    def update_document_metadata(
        self,
        document_id: str,
        user_id: str,
        title: str | None = None,
        description: str | None = None,
        author: str | None = None,
        source: str | None = None,
        year: int | None = None,
        language: str | None = None,
        keywords: List[str] | None = None,
        tags: List[str] | None = None,
    ) -> Document | None:
        """Update document metadata."""
        document = self.get_document(document_id, user_id)
        if not document:
            return None

        if title is not None:
            document.title = title
        if description is not None:
            document.description = description
        if author is not None:
            document.author = author
        if source is not None:
            document.source = source
        if year is not None:
            document.year = year
        if language is not None:
            document.language = language
        if keywords is not None:
            document.keywords = keywords

        # Update tags
        if tags is not None:
            # Remove all existing tags
            document.tags.clear()
            # Add new tags
            for tag_name in tags:
                document.add_tag(tag_name, self.db)

        document.updated_at = datetime.utcnow()
        self.db.commit()
        self.db.refresh(document)

        return document

    def delete_document(self, document_id: str, user_id: str) -> bool:
        """Delete a document and its chunks."""
        try:
            document = self.get_document(document_id, user_id)
            if not document:
                return False

            # Delete from Weaviate
            self.weaviate_service.delete_document_chunks(str(document.id))

            # Delete file from storage
            if os.path.exists(document.file_path):
                os.remove(document.file_path)

            # Delete from database (cascade will delete chunks)
            self.db.delete(document)
            self.db.commit()

            logger.info(f"Deleted document {document_id}")
            return True

        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            self.db.rollback()
            return False

    async def search_documents(
        self,
        query: str,
        user_id: str,
        limit: int = 10,
        filters: dict[str, Any] | None = None,
    ) -> list[dict[str, Any]]:
        """Search documents using semantic search with enhanced filtering."""
        try:
            # Generate query embedding using embedding service
            query_embedding = await embedding_service.generate_single_embedding(query)
            if not query_embedding:
                return []

            # Prepare filters for Weaviate
            weaviate_filters = {"user_id": user_id}
            if filters:
                # Map structured filters to Weaviate format
                if "document_type" in filters:
                    weaviate_filters["document_type"] = filters["document_type"]
                if "author" in filters:
                    weaviate_filters["author"] = filters["author"]
                if "language" in filters:
                    weaviate_filters["language"] = filters["language"]
                if "year" in filters:
                    weaviate_filters["year"] = filters["year"]
                if "chunk_type" in filters:
                    weaviate_filters["chunk_type"] = filters["chunk_type"]

            # Search in Weaviate
            search_results = self.weaviate_service.search_documents(
                query_embedding=query_embedding,
                user_id=user_id,
                limit=limit,
                filters=weaviate_filters,
            )

            # Log search query
            search_query = SearchQuery(
                user_id=uuid.UUID(user_id),
                query=query,
                query_type="knowledge",
                filters=filters or {},
                result_count=len(search_results),
            )
            self.db.add(search_query)
            self.db.commit()

            return search_results

        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []

    async def search_conversations(
        self,
        query: str,
        user_id: str,
        conversation_id: str | None = None,
        limit: int = 10,
    ) -> list[dict[str, Any]]:
        """Search conversations using semantic search."""
        try:
            # Generate query embedding using embedding service
            query_embedding = await embedding_service.generate_single_embedding(query)
            if not query_embedding:
                return []

            # Search in Weaviate
            search_results = self.weaviate_service.search_conversations(
                query_embedding=query_embedding,
                user_id=user_id,
                conversation_id=conversation_id,
                limit=limit,
            )

            # Log search query
            search_query = SearchQuery(
                user_id=uuid.UUID(user_id),
                query=query,
                query_type="conversation",
                filters={"conversation_id": conversation_id} if conversation_id else {},
                result_count=len(search_results),
            )
            self.db.add(search_query)
            self.db.commit()

            return search_results

        except Exception as e:
            logger.error(f"Error searching conversations: {e}")
            return []

    def get_search_history(
        self,
        user_id: str,
        skip: int = 0,
        limit: int = 50,
    ) -> List[SearchQuery]:
        """Get search history for a user."""
        return (
            self.db.query(SearchQuery)
            .filter(SearchQuery.user_id == user_id)
            .order_by(SearchQuery.created_at.desc())
            .offset(skip)
            .limit(limit)
            .all()
        )

    def get_tags(self, user_id: str, limit: int = 100) -> List[Tag]:
        """Get all tags for a user's documents."""
        return self.tag_service.get_tags(user_id, limit)

    def search_tags(self, query: str, user_id: str, limit: int = 20) -> List[Tag]:
        """Search tags by name."""
        return self.tag_service.search_tags(query, user_id, limit)

    def create_processing_job(
        self,
        document_id: str,
        user_id: str,
        job_type: str = "process",
        priority: int = 0,
        processing_options: Dict[str, Any] | None = None,
    ) -> DocumentProcessingJob:
        """Create a new document processing job."""
        job = DocumentProcessingJob(
            document_id=uuid.UUID(document_id),
            user_id=uuid.UUID(user_id),
            job_type=job_type,
            priority=priority,
            processing_options=processing_options or {},
        )
        
        self.db.add(job)
        self.db.commit()
        self.db.refresh(job)
        
        return job

    def get_processing_jobs(
        self,
        user_id: str,
        status: str | None = None,
        limit: int = 50,
    ) -> List[DocumentProcessingJob]:
        """Get processing jobs for a user."""
        query = self.db.query(DocumentProcessingJob).filter(
            DocumentProcessingJob.user_id == user_id
        )
        
        if status:
            query = query.filter(DocumentProcessingJob.status == status)
        
        return query.order_by(DocumentProcessingJob.priority.desc(), DocumentProcessingJob.created_at.desc()).limit(limit).all()

    def _extract_text_from_file(self, file_path: str, file_type: str) -> str | None:
        """Extract text content from various file types."""
        try:
            if file_type in ("txt", "md"):
                with open(file_path, encoding="utf-8") as f:
                    return f.read()

            elif file_type == "pdf":
                # Implement PDF text extraction
                try:
                    import io

                    import PyPDF2

                    with open(file_path, "rb") as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                        return text.strip()
                except ImportError:
                    logger.error(
                        "PyPDF2 not installed. Install with: pip install PyPDF2",
                    )
                    return None
                except Exception as e:
                    logger.error(f"PDF extraction error: {e}")
                    return None

            elif file_type in ["doc", "docx"]:
                # Implement Word document text extraction
                try:
                    import docx
                    from docx import Document

                    doc = Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"

                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                            text += "\n"

                    return text.strip()
                except ImportError:
                    logger.error(
                        "python-docx not installed. Install with: pip install python-docx",
                    )
                    return None
                except Exception as e:
                    logger.error(f"Word document extraction error: {e}")
                    return None

            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None

    def _create_chunks(
        self,
        text: str,
        document_id: str,
        chunk_size: int = 500,
        overlap: int = 50,
    ) -> list[DocumentChunk]:
        """Create text chunks from document content."""
        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for i in range(end, max(start + chunk_size - 100, start), -1):
                    if text[i] in ".!?":
                        end = i + 1
                        break

            chunk_text = text[start:end].strip()
            if chunk_text:
                # Count tokens (rough estimation)
                token_count = len(chunk_text.split())

                chunk = DocumentChunk(
                    document_id=uuid.UUID(document_id),
                    content=chunk_text,
                    chunk_index=len(chunks),
                    chunk_size=len(chunk_text),
                    token_count=token_count,
                    chunk_metadata={
                        "start_char": start,
                        "end_char": end,
                    },
                )

                self.db.add(chunk)
                chunks.append(chunk)

            start = end - overlap
            if start >= len(text):
                break

        self.db.commit()
        return chunks

    def _generate_embedding(self, text: str) -> list[float] | None:
        """Generate embedding for text using AI service."""
        try:
            # Use AI service to generate embedding
            return self.ai_service.generate_embedding(text)
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            return None
