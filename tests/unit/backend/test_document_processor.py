"""
Unit tests for Document Processor Service.

This module tests the document processor functionality including:
- File type detection
- Text extraction from different formats
- Chunking strategies
- Metadata extraction
- Document processing pipeline
"""

import io
import tempfile
from pathlib import Path
from unittest.mock import Mock, patch

import pytest
from docx import Document as DocxDocument

from backend.app.services.document import DocumentService


class TestDocumentProcessor:
    """Test class for DocumentService."""

    @pytest.fixture
    def processor(self):
        """Create DocumentService instance for testing."""
        return DocumentService()

    @pytest.fixture
    def sample_text(self):
        """Sample text for testing."""
        return """
        This is a sample document for testing purposes.
        It contains multiple paragraphs with different content.
        
        The second paragraph discusses various topics including:
        - Artificial Intelligence
        - Machine Learning
        - Natural Language Processing
        
        This document will be used to test chunking strategies and text processing.
        """

    @pytest.fixture
    def sample_pdf_content(self):
        """Sample PDF content for testing."""
        return b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n"

    @pytest.fixture
    def sample_docx_content(self):
        """Sample DOCX content for testing."""
        doc = DocxDocument()
        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("It contains multiple paragraphs.")

        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        return docx_bytes.getvalue()

    @pytest.fixture
    def sample_txt_content(self):
        """Sample TXT content for testing."""
        return b"This is a plain text document.\nIt contains multiple lines.\n"

    @pytest.fixture
    def sample_markdown_content(self):
        """Sample Markdown content for testing."""
        return b"""# Test Document

This is a **markdown** document with *formatting*.

## Section 1
- Item 1
- Item 2

## Section 2
Some content here.
"""

    def test_processor_initialization(self, processor):
        """Test DocumentProcessor initialization."""
        assert processor.chunk_size > 0
        assert processor.chunk_overlap >= 0
        assert processor.min_chunk_size > 0
        assert processor.max_chunk_size > processor.min_chunk_size
        assert processor.tokenizer is not None

    def test_detect_file_type_pdf(self, processor, sample_pdf_content):
        """Test PDF file type detection."""
        file_type = processor.detect_file_type(sample_pdf_content, "test.pdf")
        assert file_type == "application/pdf"

    def test_detect_file_type_docx(self, processor, sample_docx_content):
        """Test DOCX file type detection."""
        file_type = processor.detect_file_type(sample_docx_content, "test.docx")
        assert (
            file_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_detect_file_type_txt(self, processor, sample_txt_content):
        """Test TXT file type detection."""
        file_type = processor.detect_file_type(sample_txt_content, "test.txt")
        assert file_type == "text/plain"

    def test_detect_file_type_markdown(self, processor, sample_markdown_content):
        """Test Markdown file type detection."""
        file_type = processor.detect_file_type(sample_markdown_content, "test.md")
        assert file_type == "text/markdown"

    def test_extract_text_pdf(self, processor, sample_pdf_content):
        """Test PDF text extraction."""
        with patch(
            "backend.app.services.document_processor.pypdf.PdfReader"
        ) as mock_pdf:
            mock_reader = Mock()
            mock_page = Mock()
            mock_page.extract_text.return_value = "Extracted PDF text"
            mock_reader.pages = [mock_page]
            mock_pdf.return_value = mock_reader

            text = processor._extract_pdf_text(sample_pdf_content)
            assert text == "Extracted PDF text"

    def test_extract_text_docx(self, processor, sample_docx_content):
        """Test DOCX text extraction."""
        text = processor._extract_docx_text(sample_docx_content)
        assert "test document" in text.lower()
        assert "multiple paragraphs" in text.lower()

    def test_extract_text_txt(self, processor, sample_txt_content):
        """Test TXT text extraction."""
        text = processor._extract_txt_text(sample_txt_content)
        assert "plain text document" in text
        assert "multiple lines" in text

    def test_extract_text_markdown(self, processor, sample_markdown_content):
        """Test Markdown text extraction."""
        text = processor._extract_markdown_text(sample_markdown_content)
        assert "Test Document" in text
        assert "markdown" in text.lower()
        assert "formatting" in text.lower()

    def test_extract_text_html(self, processor):
        """Test HTML text extraction."""
        html_content = b"<html><body><h1>Title</h1><p>This is <strong>bold</strong> text.</p></body></html>"
        text = processor._extract_html_text(html_content)
        assert "Title" in text
        assert "bold" in text

    def test_extract_text_unsupported_type(self, processor):
        """Test text extraction for unsupported file type."""
        unsupported_content = b"some binary content"
        with pytest.raises(ValueError):
            processor.extract_text(unsupported_content, "test.xyz")

    def test_chunk_text_semantic(self, processor, sample_text):
        """Test semantic chunking strategy."""
        chunks = processor.chunk_text(sample_text, strategy="semantic")

        assert len(chunks) > 0
        assert all(isinstance(chunk, DocumentChunk) for chunk in chunks)
        assert all(chunk.content for chunk in chunks)
        assert all(chunk.chunk_id for chunk in chunks)

    def test_chunk_text_fixed(self, processor, sample_text):
        """Test fixed chunking strategy."""
        chunks = processor.chunk_text(sample_text, strategy="fixed", chunk_size=100)

        assert len(chunks) > 0
        assert all(isinstance(chunk, DocumentChunk) for chunk in chunks)
        assert all(len(chunk.content) <= 100 for chunk in chunks)

    def test_chunk_text_paragraph(self, processor, sample_text):
        """Test paragraph chunking strategy."""
        chunks = processor.chunk_text(sample_text, strategy="paragraph")

        assert len(chunks) > 0
        assert all(isinstance(chunk, DocumentChunk) for chunk in chunks)

    def test_chunk_text_sentence(self, processor, sample_text):
        """Test sentence chunking strategy."""
        chunks = processor.chunk_text(sample_text, strategy="sentence")

        assert len(chunks) > 0
        assert all(isinstance(chunk, DocumentChunk) for chunk in chunks)

    def test_chunk_text_invalid_strategy(self, processor, sample_text):
        """Test chunking with invalid strategy."""
        with pytest.raises(ValueError):
            processor.chunk_text(sample_text, strategy="invalid")

    def test_chunk_text_with_overlap(self, processor, sample_text):
        """Test chunking with overlap."""
        chunks = processor.chunk_text(
            sample_text, strategy="fixed", chunk_size=50, overlap=10
        )

        assert len(chunks) > 0
        # Check that consecutive chunks have some overlap
        for i in range(len(chunks) - 1):
            current_chunk = chunks[i].content
            next_chunk = chunks[i + 1].content
            # There should be some overlap in content
            assert len(set(current_chunk.split()) & set(next_chunk.split())) > 0

    def test_count_tokens(self, processor):
        """Test token counting."""
        text = "This is a test sentence."
        token_count = processor._count_tokens(text)
        assert token_count > 0
        assert isinstance(token_count, int)

    def test_calculate_importance(self, processor):
        """Test importance score calculation."""
        text = "This is an important document about AI and machine learning."
        importance = processor._calculate_importance(text)
        assert 0.0 <= importance <= 1.0

    def test_detect_language(self, processor):
        """Test language detection."""
        english_text = "This is English text."
        german_text = "Das ist deutscher Text."

        lang_en = processor._detect_language(english_text)
        lang_de = processor._detect_language(german_text)

        assert lang_en in ["en", "unknown"]
        assert lang_de in ["de", "unknown"]

    def test_extract_entities(self, processor):
        """Test entity extraction."""
        text = "Apple Inc. and Microsoft Corporation are technology companies."
        entities = processor._extract_entities(text)
        assert isinstance(entities, list)

    def test_create_enhanced_metadata(self, processor, sample_text):
        """Test enhanced metadata creation."""
        file_type = "text/plain"
        file_content = b"test content"
        chunks = [
            DocumentChunk(
                content="chunk 1",
                chunk_id="1",
                start_position=0,
                end_position=10,
                token_count=5,
                word_count=2,
            )
        ]

        metadata = processor._create_enhanced_metadata(
            file_type, file_content, sample_text, chunks
        )

        assert metadata["file_type"] == file_type
        assert metadata["file_size"] == len(file_content)
        assert metadata["text_length"] == len(sample_text)
        assert metadata["chunk_count"] == 1
        assert metadata["processing_success"] is True

    def test_extract_keywords(self, processor, sample_text):
        """Test keyword extraction."""
        keywords = processor._extract_keywords(sample_text)
        assert isinstance(keywords, list)
        assert len(keywords) > 0

    def test_extract_topics(self, processor):
        """Test topic extraction."""
        chunks = [
            DocumentChunk(
                content="This is about artificial intelligence and machine learning.",
                chunk_id="1",
                start_position=0,
                end_position=50,
                token_count=10,
                word_count=8,
            )
        ]

        topics = processor._extract_topics(chunks)
        assert isinstance(topics, list)

    def test_process_document_pdf(self, processor, sample_pdf_content):
        """Test complete document processing for PDF."""
        with patch.object(processor, "_extract_pdf_text") as mock_extract:
            mock_extract.return_value = "Extracted text from PDF"

            with patch.object(processor, "chunk_text") as mock_chunk:
                mock_chunk.return_value = [
                    DocumentChunk(
                        content="chunk 1",
                        chunk_id="1",
                        start_position=0,
                        end_position=10,
                        token_count=5,
                        word_count=2,
                    )
                ]

                result = processor.process_document(sample_pdf_content, "test.pdf")

                assert "text" in result
                assert "chunks" in result
                assert "metadata" in result
                assert result["text"] == "Extracted text from PDF"

    def test_process_document_docx(self, processor, sample_docx_content):
        """Test complete document processing for DOCX."""
        with patch.object(processor, "chunk_text") as mock_chunk:
            mock_chunk.return_value = [
                DocumentChunk(
                    content="chunk 1",
                    chunk_id="1",
                    start_position=0,
                    end_position=10,
                    token_count=5,
                    word_count=2,
                )
            ]

            result = processor.process_document(sample_docx_content, "test.docx")

            assert "text" in result
            assert "chunks" in result
            assert "metadata" in result

    def test_process_document_txt(self, processor, sample_txt_content):
        """Test complete document processing for TXT."""
        with patch.object(processor, "chunk_text") as mock_chunk:
            mock_chunk.return_value = [
                DocumentChunk(
                    content="chunk 1",
                    chunk_id="1",
                    start_position=0,
                    end_position=10,
                    token_count=5,
                    word_count=2,
                )
            ]

            result = processor.process_document(sample_txt_content, "test.txt")

            assert "text" in result
            assert "chunks" in result
            assert "metadata" in result

    def test_process_document_markdown(self, processor, sample_markdown_content):
        """Test complete document processing for Markdown."""
        with patch.object(processor, "chunk_text") as mock_chunk:
            mock_chunk.return_value = [
                DocumentChunk(
                    content="chunk 1",
                    chunk_id="1",
                    start_position=0,
                    end_position=10,
                    token_count=5,
                    word_count=2,
                )
            ]

            result = processor.process_document(sample_markdown_content, "test.md")

            assert "text" in result
            assert "chunks" in result
            assert "metadata" in result

    def test_process_document_unsupported_type(self, processor):
        """Test document processing for unsupported file type."""
        unsupported_content = b"some binary content"

        with pytest.raises(ValueError):
            processor.process_document(unsupported_content, "test.xyz")

    def test_chunk_size_validation(self, processor, sample_text):
        """Test chunk size validation."""
        # Test minimum chunk size
        with pytest.raises(ValueError):
            processor.chunk_text(sample_text, strategy="fixed", chunk_size=10)

        # Test maximum chunk size
        with pytest.raises(ValueError):
            processor.chunk_text(sample_text, strategy="fixed", chunk_size=2000)

    def test_overlap_validation(self, processor, sample_text):
        """Test overlap validation."""
        # Test negative overlap
        with pytest.raises(ValueError):
            processor.chunk_text(
                sample_text, strategy="fixed", chunk_size=100, overlap=-10
            )

        # Test overlap larger than chunk size
        with pytest.raises(ValueError):
            processor.chunk_text(
                sample_text, strategy="fixed", chunk_size=100, overlap=150
            )

    def test_document_chunk_creation(self, processor):
        """Test DocumentChunk creation and properties."""
        chunk = DocumentChunk(
            content="Test content",
            chunk_id="test_1",
            start_position=0,
            end_position=12,
            token_count=3,
            word_count=2,
            chunk_type="text",
            metadata={"source": "test"},
            importance_score=0.8,
            language="en",
            entities=["test"],
        )

        assert chunk.content == "Test content"
        assert chunk.chunk_id == "test_1"
        assert chunk.token_count == 3
        assert chunk.word_count == 2
        assert chunk.importance_score == 0.8
        assert chunk.language == "en"

    def test_document_metadata_creation(self, processor):
        """Test DocumentMetadata creation and properties."""
        metadata = DocumentMetadata(
            file_type="text/plain",
            file_size=1024,
            text_length=500,
            word_count=100,
            chunk_count=5,
            processing_engine="test_engine",
            processing_success=True,
            language="en",
            title="Test Document",
            author="Test Author",
            keywords=["test", "document"],
            summary="A test document",
            reading_time_minutes=2.5,
            complexity_score=0.7,
            topics=["technology", "testing"],
            entities=["test", "document"],
        )

        assert metadata.file_type == "text/plain"
        assert metadata.file_size == 1024
        assert metadata.text_length == 500
        assert metadata.word_count == 100
        assert metadata.chunk_count == 5
        assert metadata.processing_success is True
        assert metadata.language == "en"
        assert metadata.title == "Test Document"
        assert metadata.author == "Test Author"

    def test_calculate_avg_sentence_length(self, processor):
        """Test average sentence length calculation."""
        text = "This is sentence one. This is sentence two. This is sentence three."
        avg_length = processor._calculate_avg_sentence_length(text)
        assert avg_length > 0
        assert isinstance(avg_length, float)

    def test_get_overlap_text(self, processor):
        """Test overlap text extraction."""
        text = "This is a test sentence with multiple words for overlap testing."
        overlap_text = processor._get_overlap_text(text, 3)
        assert len(overlap_text.split()) <= 3

    def test_process_pdf_file(self, processor):
        """Test processing PDF file from path."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_file:
            temp_file.write(
                b"%PDF-1.4\n1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n"
            )
            temp_file_path = temp_file.name

        try:
            with patch.object(processor, "_extract_pdf_text") as mock_extract:
                mock_extract.return_value = "Extracted PDF text"
                result = processor.process_pdf(temp_file_path)
                assert result == "Extracted PDF text"
        finally:
            Path(temp_file_path).unlink()

    def test_process_docx_file(self, processor):
        """Test processing DOCX file from path."""
        doc = DocxDocument()
        doc.add_paragraph("Test document content.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            doc.save(temp_file.name)
            temp_file_path = temp_file.name

        try:
            result = processor.process_docx(temp_file_path)
            assert "Test document content" in result
        finally:
            Path(temp_file_path).unlink()

    def test_process_txt_file(self, processor):
        """Test processing TXT file from path."""
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as temp_file:
            temp_file.write(b"Test text content.")
            temp_file_path = temp_file.name

        try:
            result = processor.process_txt(temp_file_path)
            assert "Test text content" in result
        finally:
            Path(temp_file_path).unlink()

    def test_error_handling_file_not_found(self, processor):
        """Test error handling for non-existent files."""
        with pytest.raises(FileNotFoundError):
            processor.process_pdf("non_existent.pdf")

    def test_error_handling_corrupted_pdf(self, processor):
        """Test error handling for corrupted PDF files."""
        corrupted_pdf = b"This is not a valid PDF file"

        with pytest.raises(Exception):
            processor._extract_pdf_text(corrupted_pdf)

    def test_error_handling_corrupted_docx(self, processor):
        """Test error handling for corrupted DOCX files."""
        corrupted_docx = b"This is not a valid DOCX file"

        with pytest.raises(Exception):
            processor._extract_docx_text(corrupted_docx)

    def test_chunking_strategies_comparison(self, processor, sample_text):
        """Test different chunking strategies produce different results."""
        semantic_chunks = processor.chunk_text(sample_text, strategy="semantic")
        fixed_chunks = processor.chunk_text(
            sample_text, strategy="fixed", chunk_size=100
        )
        paragraph_chunks = processor.chunk_text(sample_text, strategy="paragraph")
        sentence_chunks = processor.chunk_text(sample_text, strategy="sentence")

        # All strategies should produce chunks
        assert len(semantic_chunks) > 0
        assert len(fixed_chunks) > 0
        assert len(paragraph_chunks) > 0
        assert len(sentence_chunks) > 0

        # Different strategies should produce different numbers of chunks
        chunk_counts = [
            len(semantic_chunks),
            len(fixed_chunks),
            len(paragraph_chunks),
            len(sentence_chunks),
        ]
        assert len(set(chunk_counts)) > 1  # At least some strategies should differ

    def test_metadata_extraction_completeness(self, processor, sample_text):
        """Test that metadata extraction includes all required fields."""
        file_type = "text/plain"
        file_content = b"test content"
        chunks = [
            DocumentChunk(
                content="chunk 1",
                chunk_id="1",
                start_position=0,
                end_position=10,
                token_count=5,
                word_count=2,
            )
        ]

        metadata = processor._create_enhanced_metadata(
            file_type, file_content, sample_text, chunks
        )

        required_fields = [
            "file_type",
            "file_size",
            "text_length",
            "word_count",
            "chunk_count",
            "processing_engine",
            "processing_success",
            "language",
        ]

        for field in required_fields:
            assert field in metadata
